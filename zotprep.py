#!/usr/bin/env python3
"""
zotprep.py — Turn a manuscript (.docx/.txt) with FAKE numbered citations
(superscript runs or [1]/(1) text) + a plain-text bibliography into a
Zotero-ready package.

v3: VERIFIED MATCHING. Earlier versions took CrossRef's #1 search result on
faith — for common surnames / generic titles / big collaborative papers
(GBD-type citations especially) that #1 result is often the WRONG paper.
This version pulls several candidates per reference (CrossRef + PubMed),
scores each against your actual reference text (author-lastname match +
year match + fuzzy title/text similarity), and only auto-accepts a match
above a confidence threshold. Anything below threshold is NOT pushed to
Zotero and NOT silently cited — it's flagged so you fix it by hand.

TWO MODES
---------
Default (static, no plugin, no account):
  Emits {Author, Year} RTF-Scan markers + a DOI list for confident matches.
  Low-confidence refs get a {NEEDS REVIEW: n} placeholder instead of a wrong
  citation. You bulk-import the DOIs with Zotero's magic wand, save as .rtf,
  run Tools > RTF Scan.

--live (live, restyle-able citations):
  Only creates Zotero items for matches at/above the confidence threshold.
  Writes Scannable Cite markers for those. Low-confidence refs are left as
  {NEEDS REVIEW: n} in the docx (not linked to any item) so you never get a
  live citation pointing at the wrong paper. Needs --zotero-userid and
  --zotero-key (https://www.zotero.org/settings/keys).

Add --mock to run fully offline (fakes CrossRef/PubMed/Zotero) for a demo.
"""
import argparse, csv, difflib, json, os, re, sys, urllib.parse, urllib.request

SUP = str.maketrans("\u2070\u00b9\u00b2\u00b3\u2074\u2075\u2076\u2077\u2078\u2079", "0123456789")
HEADERS = re.compile(r"^\s*(references|bibliography|works cited|reference list)\s*:?\s*$", re.I)
YEAR = re.compile(r"\b(19|20)\d{2}\b")
GROUP = re.compile(r"[\[(]\s*([\d\s,\-\u2013\u2014]+?)\s*[\])]")
SUPRUN = re.compile(r"[\u2070\u00b9\u00b2\u00b3\u2074\u2075\u2076\u2077\u2078\u2079]+")

CONFIDENCE_THRESHOLD = 0.55   # below this -> NEEDS REVIEW, nothing auto-created

# ---------- bibliography parsing ----------
def parse_bibliography(text):
    entries, cur, num = {}, [], None
    lead = re.compile(r"^\s*[\[(]?(\d{1,3})[\])\.\)]\s+(.*)$")
    for line in text.splitlines():
        m = lead.match(line)
        if m:
            if num is not None:
                entries[num] = " ".join(cur).strip()
            num, cur = int(m.group(1)), [m.group(2)]
        elif line.strip() and num is not None:
            cur.append(line.strip())
    if num is not None:
        entries[num] = " ".join(cur).strip()
    if not entries:
        for i, line in enumerate((l for l in text.splitlines() if l.strip()), 1):
            entries[i] = line.strip()
    return entries

def parse_ref_hint(ref):
    """Pull a rough first-author lastname + year straight from the raw
    reference text, used as the ground truth to score candidates against."""
    y = YEAR.search(ref)
    a = re.match(r"\s*([A-Z][A-Za-z'\-]+)", ref)
    return {"author": a.group(1) if a else None, "year": int(y.group()) if y else None}

# ---------- CrossRef ----------
def crossref_candidates(ref, mailto, rows=5):
    q = urllib.parse.urlencode({"query.bibliographic": ref, "rows": rows})
    url = f"https://api.crossref.org/works?{q}"
    req = urllib.request.Request(url, headers={"User-Agent": f"zotprep/1.0 (mailto:{mailto})"})
    with urllib.request.urlopen(req, timeout=20) as r:
        return json.load(r)["message"]["items"]

def cr_light(cr):
    au = (cr.get("author") or [{}])[0]
    dp = (cr.get("issued", {}).get("date-parts", [[None]])[0] or [None])
    return {"source": "crossref", "author": au.get("family") or au.get("name") or "",
            "year": dp[0], "title": (cr.get("title") or [""])[0],
            "doi": cr.get("DOI"), "raw": cr}

# ---------- PubMed (better hit rate for medical/epi references) ----------
def pubmed_candidates(ref, mailto, rows=3):
    base = "https://eutils.ncbi.nlm.nih.gov/entrez/eutils"
    q = urllib.parse.urlencode({"db": "pubmed", "retmax": rows, "term": ref,
                                 "retmode": "json", "tool": "zotprep", "email": mailto})
    with urllib.request.urlopen(f"{base}/esearch.fcgi?{q}", timeout=20) as r:
        ids = json.load(r)["esearchresult"].get("idlist", [])
    if not ids:
        return []
    q2 = urllib.parse.urlencode({"db": "pubmed", "id": ",".join(ids), "retmode": "json",
                                  "tool": "zotprep", "email": mailto})
    with urllib.request.urlopen(f"{base}/esummary.fcgi?{q2}", timeout=20) as r:
        summ = json.load(r)["result"]
    out = []
    for pmid in ids:
        rec = summ.get(pmid)
        if not rec:
            continue
        au = (rec.get("authors") or [{}])[0].get("name", "")
        y = YEAR.search(rec.get("pubdate", ""))
        doi = next((idr["value"] for idr in rec.get("articleids", [])
                    if idr.get("idtype") == "doi"), None)
        out.append({"source": "pubmed", "author": au.split()[0] if au else "",
                     "year": int(y.group()) if y else None, "title": rec.get("title", ""),
                     "doi": doi, "raw": rec})
    return out

# ---------- mock candidates (offline demo) ----------
def mock_candidates(ref, n):
    parts = [p.strip() for p in ref.split(". ") if p.strip()]
    last = re.split(r"[ ,]", parts[0])[0] if parts else "Unknown"
    title = parts[1] if len(parts) > 1 else ""
    y = YEAR.search(ref)
    good = {"source": "crossref", "author": last, "year": int(y.group()) if y else None,
            "title": title, "doi": f"10.1000/mock.{n}", "raw": {}}
    return [good]

# ---------- scoring ----------
def score_candidate(ref_text, hint, cand):
    """0..1 confidence that `cand` is the paper `ref_text` actually cites.
    Hard signal: does the candidate's author lastname appear in the raw
    reference, and does the year match exactly? Soft signal: fuzzy text
    similarity between the raw reference and a reconstructed cand string."""
    author_ok = bool(cand["author"]) and cand["author"].lower() in ref_text.lower()
    year_ok = hint["year"] is not None and cand["year"] == hint["year"]
    recon = f'{cand["author"]} {cand["year"]} {cand["title"]}'.strip()
    fuzz = difflib.SequenceMatcher(None, recon.lower(), ref_text.lower()).ratio()
    score = fuzz
    score += 0.25 if author_ok else -0.20
    score += 0.15 if year_ok else -0.15
    return max(0.0, min(1.0, score))

def best_match(ref, hint, candidates):
    scored = [(score_candidate(ref, hint, c), c) for c in candidates]
    scored.sort(key=lambda t: t[0], reverse=True)
    return scored[0] if scored else (0.0, None)

def resolve_reference(ref, n, args):
    hint = parse_ref_hint(ref)
    if args.mock:
        cands = mock_candidates(ref, n)
    else:
        cands = []
        try:
            cands += [cr_light(c) for c in crossref_candidates(ref, args.mailto)]
        except Exception as e:
            print(f"  [{n}] CrossRef error: {e}", file=sys.stderr)
        try:
            cands += pubmed_candidates(ref, args.mailto)
        except Exception as e:
            print(f"  [{n}] PubMed error: {e}", file=sys.stderr)
    score, cand = best_match(ref, hint, cands)
    ok = cand is not None and score >= CONFIDENCE_THRESHOLD and cand.get("doi")
    if not cand:
        cand = {"author": hint["author"] or "Anon", "year": hint["year"] or "n.d.",
                 "title": "", "doi": None, "source": "none", "raw": None}
    return {"author": cand["author"] or "Anon", "year": cand["year"] or "n.d.",
            "title": cand.get("title", ""), "doi": cand.get("doi"),
            "score": round(score, 2), "ok": ok, "raw": cand.get("raw")}

# ---------- CrossRef/PubMed raw -> Zotero item (for --live) ----------
def raw_to_item(meta):
    cr = meta.get("raw")
    if meta.get("doi") and isinstance(cr, dict) and "title" in cr and "author" in cr:
        # crossref raw record
        tmap = {"journal-article": "journalArticle", "book": "book",
                "book-chapter": "bookSection", "proceedings-article": "conferencePaper",
                "posted-content": "preprint"}
        itype = tmap.get(cr.get("type"), "journalArticle")
        creators = [{"creatorType": "author", "firstName": a.get("given", ""),
                     "lastName": a.get("family") or a.get("name", "")} for a in cr.get("author", [])]
        dp = (cr.get("issued", {}).get("date-parts", [[None]])[0] or [None])
        item = {"itemType": itype, "title": (cr.get("title") or [""])[0], "creators": creators,
                "date": str(dp[0]) if dp[0] else "", "DOI": cr.get("DOI", "")}
        ct = (cr.get("container-title") or [""])
        if itype == "journalArticle":
            item["publicationTitle"] = ct[0] if ct else ""
        for k_cr, k_zot in (("volume", "volume"), ("issue", "issue"), ("page", "pages")):
            if cr.get(k_cr):
                item[k_zot] = cr[k_cr]
        return item
    # generic fallback (pubmed / mock): minimal but correct item
    return {"itemType": "journalArticle", "title": meta.get("title", ""),
            "creators": [{"creatorType": "author", "lastName": meta.get("author", ""), "firstName": ""}],
            "date": str(meta.get("year", "")), "DOI": meta.get("doi") or ""}

def push_to_zotero(items_by_n, userid, api_key):
    from pyzotero import zotero
    zot = zotero.Zotero(userid, "user", api_key)
    keys = {}
    for start in range(0, len(items_by_n), 50):
        chunk = items_by_n[start:start + 50]
        resp = zot.create_items([it for _, it in chunk])
        for i_str, k in resp.get("success", {}).items():
            keys[chunk[int(i_str)][0]] = k
        for i_str, info in resp.get("failed", {}).items():
            print(f"  [{chunk[int(i_str)][0]}] Zotero create failed: {info}", file=sys.stderr)
    return keys

# ---------- marker rendering ----------
def make_renderer(style, meta, keys=None, uid=None):
    def expand(spec):
        out = []
        for chunk in spec.replace("\u2013", "-").replace("\u2014", "-").split(","):
            chunk = chunk.strip()
            if "-" in chunk:
                a, b = chunk.split("-", 1)
                if a.isdigit() and b.isdigit():
                    out += range(int(a), int(b) + 1)
            elif chunk.isdigit():
                out.append(int(chunk))
        return [n for n in out if n in meta]

    def one_rtf(n):
        m = meta[n]
        if not m["ok"]:
            return f"NEEDS REVIEW: ref {n}"
        if m.get("title"):
            return f'{m["author"]}, "{m["title"]}", {m["year"]}'
        return f'{m["author"]}, {m["year"]}'

    def render(spec):
        nums = expand(spec)
        if not nums:
            return None
        if style == "scannable":
            parts = []
            for n in nums:
                if meta[n]["ok"] and keys.get(n):
                    parts.append(f'{{ | {meta[n]["author"]}, ({meta[n]["year"]}) | | |zu:{uid}:{keys[n]}}}')
                else:
                    parts.append(f"{{NEEDS REVIEW: ref {n}}}")
            return "".join(parts)
        return "{" + "; ".join(one_rtf(n) for n in nums) + "}"
    return render

# ---------- apply to text / docx ----------
def rewrite_text(body, render):
    body = GROUP.sub(lambda m: render(m.group(1)) or m.group(0), body)
    body = SUPRUN.sub(lambda m: render(m.group(0).translate(SUP)) or m.group(0), body)
    return body

def to_rtf(text):
    esc = text.replace("\\", "\\\\").replace("{", "\\{").replace("}", "\\}")
    return "{\\rtf1\\ansi\\deff0\n" + esc.replace("\r\n", "\n").replace("\n", "\\par\n") + "\n}\n"

def docx_find_biblio(doc):
    for i, p in enumerate(doc.paragraphs):
        if HEADERS.match(p.text):
            return i
    return None

def docx_mark(doc, idx, render):
    for p in doc.paragraphs[:idx]:
        for r in p.runs:
            t = r.text
            if not t.strip():
                continue
            if r.font.superscript and re.search(r"\d", t):
                rep = render(t)
                if rep:
                    r.text, r.font.superscript = rep, False
                continue
            new = GROUP.sub(lambda m: render(m.group(1)) or m.group(0), t)
            if new != t:
                r.text = new

def docx_remove_from(doc, idx):
    for p in list(doc.paragraphs)[idx:]:
        p._element.getparent().remove(p._element)

# ---------- main ----------
def main():
    global CONFIDENCE_THRESHOLD
    ap = argparse.ArgumentParser(description="Prep a .docx/.txt paper for Zotero, with verified matching.")
    ap.add_argument("--manuscript", required=True, help=".docx or .txt")
    ap.add_argument("--bibliography", help="separate biblio file; else auto-split on 'References'")
    ap.add_argument("--outdir", default="zot_out")
    ap.add_argument("--mailto", default="you@example.com", help="email for CrossRef/PubMed polite pool")
    ap.add_argument("--live", action="store_true", help="push confident matches to Zotero + emit Scannable Cite markers")
    ap.add_argument("--zotero-userid", help="numeric Zotero user ID (for --live)")
    ap.add_argument("--zotero-key", help="Zotero API key with write access (for --live)")
    ap.add_argument("--threshold", type=float, default=CONFIDENCE_THRESHOLD,
                    help=f"confidence cutoff to auto-accept a match (default {CONFIDENCE_THRESHOLD})")
    ap.add_argument("--mock", action="store_true", help="offline demo (fake CrossRef/PubMed/Zotero)")
    args = ap.parse_args()
    CONFIDENCE_THRESHOLD = args.threshold
    os.makedirs(args.outdir, exist_ok=True)
    is_docx = args.manuscript.lower().endswith(".docx")

    if args.live and not is_docx:
        sys.exit("--live needs a .docx input (Scannable Cite markers convert in Word/LibreOffice).")
    if args.live and not args.mock and not (args.zotero_userid and args.zotero_key):
        sys.exit("--live needs --zotero-userid and --zotero-key (or --mock to demo).")

    # --- read body + bibliography ---
    doc = None
    if is_docx:
        from docx import Document
        doc = Document(args.manuscript)
        if args.bibliography:
            biblio = open(args.bibliography, encoding="utf-8").read(); idx = len(doc.paragraphs)
        else:
            idx = docx_find_biblio(doc)
            if idx is None:
                sys.exit("No 'References' heading found and no --bibliography given.")
            biblio = "\n".join(p.text for p in doc.paragraphs[idx + 1:])
    else:
        text = open(args.manuscript, encoding="utf-8").read()
        if args.bibliography:
            biblio, body = open(args.bibliography, encoding="utf-8").read(), text
        else:
            lines = text.splitlines()
            hi = next((i for i in range(len(lines) - 1, -1, -1) if HEADERS.match(lines[i])), None)
            if hi is None:
                sys.exit("No 'References' header found and no --bibliography given.")
            body, biblio = "\n".join(lines[:hi]), "\n".join(lines[hi + 1:])

    refs = parse_bibliography(biblio)
    print(f"Parsed {len(refs)} references. Confidence threshold: {CONFIDENCE_THRESHOLD}\n")

    meta = {}
    for n in sorted(refs):
        meta[n] = resolve_reference(refs[n], n, args)
        tag = "OK  " if meta[n]["ok"] else "REVIEW"
        print(f"  [{n}] {tag} score={meta[n]['score']:.2f}  {meta[n]['author']}, {meta[n]['year']}  "
              f"{meta[n]['doi'] or '(no confident DOI)'}")

    needs_review = [n for n in refs if not meta[n]["ok"]]

    # --- live: create items only for confident matches ---
    keys, uid = {}, args.zotero_userid or "0"
    if args.live:
        items_by_n = [(n, raw_to_item(meta[n])) for n in sorted(refs) if meta[n]["ok"]]
        if args.mock:
            keys = {n: f"MOCKKEY{n:02d}" for n, _ in items_by_n}
        else:
            print("\nCreating confident-match items in Zotero ...")
            keys = push_to_zotero(items_by_n, args.zotero_userid, args.zotero_key)
        print(f"  created {len(keys)} items ({len(needs_review)} left for manual review)")

    style = "scannable" if args.live else "rtfscan"
    render = make_renderer(style, meta, keys, uid)

    # --- write outputs ---
    if is_docx:
        docx_mark(doc, idx, render)
        body_text = "\n".join(p.text for p in doc.paragraphs[:idx])
        if not args.bibliography:
            docx_remove_from(doc, idx)
        out_docx = "manuscript_scannable.docx" if args.live else "manuscript_marked.docx"
        doc.save(os.path.join(args.outdir, out_docx))
        if not args.live:
            open(os.path.join(args.outdir, "manuscript.rtf"), "w").write(to_rtf(body_text))
    else:
        open(os.path.join(args.outdir, "manuscript.rtf"), "w").write(to_rtf(rewrite_text(body, render)))

    if not args.live:
        dois = [meta[n]["doi"] for n in sorted(refs) if meta[n]["ok"] and meta[n]["doi"]]
        open(os.path.join(args.outdir, "doi_list.txt"), "w").write("\n".join(dois) + ("\n" if dois else ""))

    with open(os.path.join(args.outdir, "report.csv"), "w", newline="") as f:
        w = csv.writer(f)
        w.writerow(["n", "status", "score", "doi", "author", "year", "zotero_key", "raw_reference"])
        for n in sorted(refs):
            m = meta[n]
            w.writerow([n, "OK" if m["ok"] else "NEEDS REVIEW", m["score"], m["doi"] or "",
                        m["author"], m["year"], keys.get(n, ""), refs[n]])

    print(f"\nWrote -> {args.outdir}/")
    if needs_review:
        print(f"\n*** {len(needs_review)} reference(s) NEED MANUAL REVIEW (below confidence "
              f"{CONFIDENCE_THRESHOLD}): {needs_review} ***")
        print("These were NOT pushed to Zotero / NOT auto-cited. Marked '{NEEDS REVIEW: n}' in the "
              "document — search and add them by hand, then replace that marker.")
    if args.live:
        print("\nNext: Tools > ODF Scan on manuscript_scannable.docx (ODF/DOCX Scan plugin).")
    else:
        print("\nNext: paste doi_list.txt into the magic wand, then Tools > RTF Scan on the .rtf.")

if __name__ == "__main__":
    main()
