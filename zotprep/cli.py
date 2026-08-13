"""Command line entry point.

    python -m zotprep --manuscript yourpaper.docx --dry-run

    # once, to stop re-entering credentials on every run:
    python -m zotprep --zotero-userid 1234567 --zotero-key XXXX --save-credentials

    python -m zotprep --manuscript yourpaper.docx --live
"""
from __future__ import annotations

import argparse
import asyncio
import csv
import os
import sys

import httpx

from . import config
from .cache import Cache
from .docx_writer import (
    biblio_end_index,
    count_citations,
    find_biblio_index,
    make_renderer,
    mark_body,
    mark_body_fields,
    parse_bibliography,
    remove_range,
)
from .extractor import parse_reference
from .resolver import USER_AGENT, resolve_all
from .zotero.client import ZoteroWriter, enrich, to_item


def build_parser() -> argparse.ArgumentParser:
    ap = argparse.ArgumentParser(prog="zotprep", description="Resolve manuscript references and prepare a Zotero-linked .docx")
    ap.add_argument("--manuscript", help=".docx containing the manuscript")
    ap.add_argument("--bibliography", help="separate bibliography file; else auto-split on 'References'")
    ap.add_argument("--outdir", default="zot_out")
    ap.add_argument("--mailto", help="contact email for the Crossref/NCBI polite pools")
    ap.add_argument("--workers", type=int, default=12)
    ap.add_argument("--no-cache", action="store_true")
    ap.add_argument("--dry-run", action="store_true",
                    help="resolve and report, but create nothing in Zotero")
    ap.add_argument("--live", action="store_true",
                    help="create/reuse items in your Zotero library (must be explicit)")
    ap.add_argument("--zotero-userid",
                    help="numeric Zotero userID; else $ZOTERO_USERID, else saved config")
    ap.add_argument("--zotero-key",
                    help="read/write Zotero API key; else $ZOTERO_KEY, else saved config")
    ap.add_argument("--review", action="store_true",
                    help="prompt for anything not auto-accepted instead of leaving it flagged")
    ap.add_argument("--style", choices=("fields", "scannable"), default="fields",
                    help="fields: live Zotero citations, ready to open in Word (default). "
                         "scannable: Scannable Cite markers for the ODF Scan plugin")

    creds = ap.add_argument_group("saved credentials")
    creds.add_argument("--save-credentials", action="store_true",
                       help=f"write the userID/key/mailto in use to {config.config_path()} "
                            "and stop asking for them")
    creds.add_argument("--forget-credentials", action="store_true",
                       help="delete the saved credentials file")
    creds.add_argument("--show-config", action="store_true",
                       help="print where each setting is coming from, then exit")
    return ap


async def run(args) -> int:
    from docx import Document

    # Writing to someone's Zotero library is the one irreversible thing this tool
    # does, so it must be asked for explicitly — never the result of a forgotten
    # flag. --dry-run stays available as the obvious safe default.
    if args.live and args.dry_run:
        sys.exit("--live and --dry-run are mutually exclusive.")
    if not args.live:
        args.dry_run = True
    if args.live and not (args.zotero_userid and args.zotero_key):
        sys.exit(
            "--live needs a Zotero userID and API key.\n"
            "  Pass --zotero-userid/--zotero-key once with --save-credentials and they are\n"
            f"  remembered in {config.config_path()}; $ZOTERO_USERID/$ZOTERO_KEY also work.\n"
            "  Get both at https://www.zotero.org/settings/keys (key needs write access)."
        )

    os.makedirs(args.outdir, exist_ok=True)

    # A mistyped filename is the most likely thing to go wrong here, and a
    # python-docx traceback is a poor way to say so. The .doc/.pdf case is worth
    # naming separately: those open as "not a Word file", which is true but
    # unhelpful when the fix is simply to re-save as .docx.
    if not os.path.isfile(args.manuscript):
        sys.exit(f"Cannot read {args.manuscript}: no such file.")
    try:
        doc = Document(args.manuscript)
    except Exception as exc:  # noqa: BLE001 - python-docx raises several types
        hint = ""
        if args.manuscript.lower().endswith((".doc", ".pdf", ".rtf", ".odt")):
            hint = ("\n  Only .docx is supported. Open it in Word and use "
                    "File > Save As > Word Document (.docx).")
        sys.exit(f"Cannot read {args.manuscript}: {exc}{hint}")

    if args.bibliography:
        try:
            biblio_text = open(args.bibliography, encoding="utf-8").read()
        except OSError as exc:
            sys.exit(f"Cannot read {args.bibliography}: {exc.strerror or exc}.")
        biblio_idx = biblio_end = len(doc.paragraphs)
    else:
        biblio_idx = find_biblio_index(doc)
        if biblio_idx is None:
            sys.exit("No 'References' heading found; pass --bibliography.")
        # Bounded, so that trailing figures and tables are neither read as
        # references nor removed with them.
        biblio_end = biblio_end_index(doc, biblio_idx)
        biblio_text = "\n".join(p.text for p in doc.paragraphs[biblio_idx + 1:biblio_end])

    raw_refs = parse_bibliography(biblio_text)
    refs = {n: parse_reference(n, t) for n, t in raw_refs.items()}
    print(f"Parsed {len(refs)} references.")

    # A bibliography with nothing citing it means the in-text notation was not
    # recognised. Stop here: the alternative is minutes of resolution followed by
    # a document whose reference list has been removed and nothing linked.
    n_found = count_citations(doc, biblio_idx)
    if not n_found:
        sys.exit(
            f"Parsed {len(refs)} references but found no citations in the body.\n"
            "  Recognised notations are [1] and (2,3), superscript digits, and plain\n"
            "  digits such as '...knee OA 2.' or '...analgesia 6,7.'\n"
            "  Nothing was changed and no items were added."
        )
    print(f"Found {n_found} in-text citation markers.")

    cache = Cache(enabled=not args.no_cache)
    done = {"n": 0}

    def progress(_res):
        done["n"] += 1
        print(f"  resolving {done['n']}/{len(refs)}", end="\r", file=sys.stderr)

    results = await resolve_all(refs, args.mailto, cache, workers=args.workers, progress=progress)
    print()

    if args.review:
        from .review import review_loop

        review_loop(refs, results, cache)

    accepted = [n for n in sorted(results) if results[n].status in ("ACCEPTED", "FROM_TEXT")]
    # A reference deleted at review is decided, not outstanding.
    dropped = [n for n in sorted(results) if results[n].status == "DROPPED"]
    unresolved = [n for n in sorted(results) if n not in accepted and n not in dropped]

    # upgrade winning candidates to canonical Crossref records for item quality
    async with httpx.AsyncClient(headers={"User-Agent": USER_AGENT}, timeout=30.0,
                                 follow_redirects=True) as client:
        for n in accepted:
            res = results[n]
            if res.candidate:
                res.candidate = await enrich(client, res.candidate, args.mailto)

    writer = ZoteroWriter(args.zotero_userid or "0", args.zotero_key or "", dry_run=args.dry_run)
    if not args.dry_run:
        print("Indexing existing Zotero library for duplicates ...")
        writer.load_library()
    items = [(n, to_item(results[n].candidate, refs[n])) for n in accepted if results[n].candidate]
    keys = writer.create(items)
    for n, k in keys.items():
        results[n].zotero_key = k
    print(f"Zotero items ready: {len(keys)} ({'dry run' if args.dry_run else 'created/reused'})")

    warnings: list[str] = []
    render = make_renderer(results, keys, args.zotero_userid or "0", refs=refs,
                           warnings=warnings, style=args.style)
    if args.style == "fields":
        n_marked = mark_body_fields(doc, biblio_idx, render)
    else:
        n_marked = mark_body(doc, biblio_idx, render)
    # The reference list goes only when something has taken its place. Removing
    # it after marking nothing would hand back a manuscript with no citations
    # and no bibliography either.
    if not args.bibliography and n_marked:
        remove_range(doc, biblio_idx, biblio_end)
    elif not args.bibliography:
        print("No citations were rewritten — the bibliography has been left in place.")
    name = "manuscript_zotero.docx" if args.style == "fields" else "manuscript_scannable.docx"
    out_docx = os.path.join(args.outdir, name)
    doc.save(out_docx)

    report = os.path.join(args.outdir, "report.csv")
    with open(report, "w", newline="", encoding="utf-8") as f:
        w = csv.writer(f)
        w.writerow(["n", "status", "tier", "confidence", "doi", "pmid", "zotero_key",
                    "resolved_title", "reason", "advisory", "raw_reference"])
        for n in sorted(results):
            r = results[n]
            c = r.candidate
            w.writerow([n, r.status, r.tier, r.confidence, (c.doi if c else ""),
                        (c.pmid if c else ""), r.zotero_key or "",
                        (c.title if c else ""), r.reason,
                        " | ".join(r.notes), refs[n].raw])

    print(f"\nWrote {out_docx}")
    print(f"Wrote {report}   ({n_marked} in-text citations rewritten)")
    advisories = [(n, note) for n in sorted(results) for note in results[n].notes]
    if advisories:
        print("\n*** Bibliography entries to verify "
              "(resolved, but the entry's own fields disagree):")
        for n, note in advisories:
            print(f"  [{n}] {note}")
    if warnings:
        print("\n*** In-text citation problems (document left unchanged at these spots):")
        for w in dict.fromkeys(warnings):
            print(f"  - {w}")
    if dropped:
        print(f"\n{len(dropped)} citation(s) deleted at review: {dropped} "
              "— the markers were removed from the text.")
    if unresolved:
        print(f"\n*** {len(unresolved)} reference(s) unresolved: {unresolved}")
        print("Marked '{NEEDS REVIEW: n}' in the document. Re-run with --review to fix interactively.")
    elif args.style == "fields":
        print("\nAll references resolved. The .docx has live Zotero citations — open it in "
              "Word;\nZotero will ask for a citation style the first time you refresh.")
    else:
        print("\nAll references resolved. Next: Zotero > Tools > ODF Scan on the output .docx")
    cache.close()
    return 1 if unresolved else 0


def apply_settings(args) -> None:
    """Fold saved config / environment into args, and service the credential
    management flags. Runs before any work so --show-config and
    --forget-credentials do not require a manuscript."""
    cli = {
        "zotero_userid": args.zotero_userid,
        "zotero_key": args.zotero_key,
        "mailto": args.mailto,
    }

    if args.forget_credentials:
        path = config.config_path()
        print(f"Removed {path}" if config.forget() else f"Nothing saved at {path}")
        if not args.manuscript:
            sys.exit(0)

    if args.show_config:
        print(f"config file: {config.config_path()}")
        resolved = config.resolve(cli)
        for field in config.FIELDS:
            value = resolved[field] or ""
            if field == "zotero_key" and value:
                value = f"{value[:4]}...{value[-4:]}"
            print(f"  {field:<14} {value or '(unset)':<28} from {config.source_of(field, cli)}")
        sys.exit(0)

    if args.save_credentials:
        saved = {k: v for k, v in config.resolve(cli).items() if v}
        if not saved:
            sys.exit("--save-credentials: nothing to save. Pass --zotero-userid/--zotero-key "
                     "(and optionally --mailto) alongside it.")
        print(f"Saved {', '.join(sorted(saved))} to {config.save(saved)}")
        if not args.manuscript:
            sys.exit(0)

    resolved = config.resolve(cli)
    args.zotero_userid = resolved["zotero_userid"]
    args.zotero_key = resolved["zotero_key"]
    args.mailto = resolved["mailto"] or "you@example.com"


def main() -> None:
    args = build_parser().parse_args()
    apply_settings(args)
    if not args.manuscript:
        build_parser().error("the following arguments are required: --manuscript")
    sys.exit(asyncio.run(run(args)))


if __name__ == "__main__":
    main()
