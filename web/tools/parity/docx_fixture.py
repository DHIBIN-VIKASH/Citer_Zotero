"""Build a .docx fixture and record what the Python docx_writer does to it.

The browser port of docx_writer.py cannot be compared by feeding it strings: its
whole difficulty is that Word splits a single citation across runs, stores
superscripts as run formatting rather than characters, and hides the text the
scanner measures behind python-docx's run/paragraph semantics. So the fixture is
a real .docx containing every shape that logic has to survive, and the
comparison is the resulting paragraph text — the thing that actually ships.

Writes:
  web/fixture.docx        the input, for the browser to fetch
  web/fixture.expect.json what Python produces from it

Both are gitignored (*.docx, and the json is written next to it). Run:
    python web/tools/parity/docx_fixture.py
then open the app's page and run the docx parity check from the browser console.
"""
from __future__ import annotations

import base64
import json
import os
import sys
import tempfile
from pathlib import Path

ROOT = Path(__file__).resolve().parents[3]
sys.path.insert(0, str(ROOT))

from docx import Document  # noqa: E402

from docx.enum.section import WD_ORIENT, WD_SECTION  # noqa: E402
from docx.shared import Inches  # noqa: E402

from zotprep.docx_writer import (  # noqa: E402
    biblio_end_index, find_biblio_index, make_renderer, mark_body, parse_bibliography,
    remove_range,
)

WEB = ROOT / "web"
SUP = "⁰¹²³⁴⁵⁶⁷⁸⁹"
W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

# 1x1 transparent PNG — the smallest thing add_picture accepts.
PNG = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mP8z8BQDwAEhQGAhKmM"
    "IQAAAABJRU5ErkJggg=="
)


def outline(doc) -> list[list]:
    """[tag, text, image count, page size] for every body child.

    The paragraph texts alone cannot show that the trailing matter survived: a
    deleted image, two tables fused into one, or a lost section break all leave
    the remaining paragraph text identical. This is what the two engines are
    compared on instead.
    """
    rows = []
    for el in doc.element.body:
        tag = el.tag.replace(W, "")
        sect = el.find(f"{W}pPr/{W}sectPr") if tag == "p" else (el if tag == "sectPr" else None)
        page = None
        if sect is not None:
            sz = sect.find(f"{W}pgSz")
            if sz is not None:
                page = [sz.get(f"{W}orient"), sz.get(f"{W}w"), sz.get(f"{W}h")]
        rows.append([
            tag,
            "".join(t.text or "" for t in el.iter(f"{W}t")),
            sum(1 for _ in el.iter(f"{W}drawing")),
            page,
        ])
    return rows


def sup(n: str) -> str:
    """ASCII digits -> Unicode superscript digits."""
    return "".join(SUP[int(c)] if c.isdigit() else c for c in n)


def build() -> Path:
    doc = Document()
    doc.add_paragraph("A Manuscript With Citations")

    # 1. Unicode superscript citations, single and ranged, incl. the Lancet
    #    middle-dot separator.
    doc.add_paragraph(f"Health systems vary widely.{sup('1')} Others disagree.{sup('2')}⁻{sup('4')}")
    doc.add_paragraph(f"Multiple, dot separated.{sup('1')}·{sup('3')}")

    # 2. Bracketed groups.
    doc.add_paragraph("Bracketed forms are common [1] and also (2,3) and [4-6].")

    # 3. Digits fused onto sentence punctuation — a degraded superscript.
    doc.add_paragraph("Better-resourced systems fare differently.5-7 And again.8")

    # 4. Mathematics that must NOT become citations.
    doc.add_paragraph(f"The rate was -13.3x10⁻{sup('3')} per year, with R{sup('2')}=0.32 and 5 km{sup('2')}.")

    # 5. A citation split across runs, which is why the scan works in paragraph
    #    coordinates rather than per-run.
    p = doc.add_paragraph("Split across runs;")
    r1 = p.add_run(sup("1"))
    r2 = p.add_run(sup("2"))
    r1.font.superscript = False  # already superscript characters
    r2.font.superscript = False
    p.add_run(" continues.")

    # 6. Word superscript *formatting* over ordinary digits.
    p = doc.add_paragraph("Formatted superscript follows this")
    rr = p.add_run("9")
    rr.font.superscript = True
    p.add_run(" and then prose.")

    # 7. An oversized range: a dropped digit, which must be refused.
    doc.add_paragraph(f"A dropped digit produces{sup('3')}⁻{sup('3')}{sup('2')} a huge span.")

    # 8. A year range that is not a citation at all.
    doc.add_paragraph("Between (1990-2023) the trend held.")

    doc.add_paragraph("References")
    doc.add_paragraph("1. Barro RJ, Sala-i-Martin X. Convergence. J Polit Econ. 1992;100(2):223-51.")
    doc.add_paragraph("2. Marmot M. Social determinants of health. Lancet. 2005;365(9464):1099-104.")
    doc.add_paragraph("3. Theil H. Economics and Information Theory. Amsterdam: North-Holland; 1967.")
    doc.add_paragraph("4. GBD 2021 Collaborators. Global incidence. Lancet. 2024.")
    doc.add_paragraph("5. Smith AB. A fifth reference. BMJ. 2010;340:c2289.")
    doc.add_paragraph("6. Jones CD. A sixth reference. Nature. 2015;520:100-9.")
    doc.add_paragraph("7. Lee K. A seventh reference. Cell. 2018;172:1-10.")
    doc.add_paragraph("8. Park S. An eighth reference. Science. 2019;363:200-5.")
    doc.add_paragraph("9. Choi H. A ninth reference. PNAS. 2020;117:5000-10.")

    # 9. Trailing matter: everything that must survive the bibliography being
    #    removed. The blank paragraph is the author's spacing, the legend carries
    #    an image, and the tables sit in a landscape section whose break lives in
    #    a paragraph mark — the one that turns the whole document landscape if it
    #    is deleted with the reference list.
    doc.add_paragraph("")
    doc.add_paragraph("Figure 1. Study selection flow diagram.")
    with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as fh:
        fh.write(PNG)
        png = fh.name
    try:
        doc.add_picture(png, width=Inches(1))
    finally:
        os.unlink(png)

    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = Inches(11.69), Inches(8.27)
    doc.add_paragraph("Tables")
    doc.add_paragraph("Table 1. Characteristics of the included studies.")
    doc.add_table(rows=2, cols=3)
    doc.add_paragraph("NR = not reported.")
    doc.add_paragraph("Table 2. Pooled estimates by subgroup.")
    doc.add_table(rows=2, cols=2)

    out = WEB / "fixture.docx"
    doc.save(out)
    return out


class StubCandidate:
    def __init__(self, n):
        self.year = 1990 + n
        self.corporate = "GBD 2021 Collaborators" if n == 4 else None
        self.authors = [f"surname{n}"]


class StubResolution:
    def __init__(self, n, status):
        self.status = status
        self.candidate = StubCandidate(n) if status != "REVIEW" else None


def main() -> int:
    path = build()

    doc = Document(path)
    idx = find_biblio_index(doc)
    end = biblio_end_index(doc, idx)
    biblio = "\n".join(p.text for p in doc.paragraphs[idx + 1:end])
    raw = parse_bibliography(biblio)

    # Reference 6 is deliberately left unresolved, so the "{NEEDS REVIEW}" branch
    # is exercised alongside the live-marker one.
    resolutions = {n: StubResolution(n, "REVIEW" if n == 6 else "ACCEPTED") for n in raw}
    keys = {n: f"KEY{n:03d}" for n in raw if n != 6}

    warnings: list[str] = []
    render = make_renderer(resolutions, keys, "1234567", refs=None, warnings=warnings)
    n_marked = mark_body(doc, idx, render)
    remove_range(doc, idx, end)

    expect = {
        "biblio_index": idx,
        "biblio_end": end,
        "entries": {str(k): v for k, v in raw.items()},
        "paragraphs": [p.text for p in doc.paragraphs],
        "outline": outline(doc),
        "n_marked": n_marked,
        "warnings": sorted(set(warnings)),
    }
    (WEB / "fixture.expect.json").write_text(
        json.dumps(expect, ensure_ascii=False, indent=1), encoding="utf-8"
    )
    print(f"wrote {path.name} and fixture.expect.json")
    print(f"  bibliography at paragraphs [{idx}, {end}), {len(raw)} entries, {n_marked} markers")
    kept = expect["outline"]
    print(f"  kept after removal: {sum(1 for r in kept if r[0] == 'tbl')} tables, "
          f"{sum(r[2] for r in kept)} images, {sum(1 for r in kept if r[3])} section breaks")
    for w in expect["warnings"]:
        print(f"  warning: {w}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
