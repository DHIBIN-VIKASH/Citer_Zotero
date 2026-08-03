"""Check that removing the bibliography removes nothing else.

A manuscript keeps things after its reference list: figure legends and the
images they caption, table captions and the tables themselves, appendices — and,
invisibly, the section break that ends the portrait section before a landscape
table page. Deleting every paragraph from the "References" heading to the end of
the document takes all of that with it, and the section-break loss is the worst
of the three: with the break gone, everything before it joins the following
section, so one landscape table at the end turns the whole document landscape.

Each case below builds a document, performs the surgery the CLI performs, and
asserts that the parts outside the bibliography came through byte-for-byte:
same body children in the same order, same text, same images, same section
properties.

    python tests/docx_surgery_check.py
"""
from __future__ import annotations

import base64
import io
import os
import sys
import tempfile

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
sys.stdout.reconfigure(encoding="utf-8")

from docx import Document  # noqa: E402
from docx.enum.section import WD_ORIENT, WD_SECTION  # noqa: E402
from docx.shared import Inches  # noqa: E402

from zotprep.docx_writer import (  # noqa: E402
    biblio_end_index,
    find_biblio_index,
    parse_bibliography,
    remove_range,
)

W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

# 1x1 transparent PNG — the smallest thing add_picture accepts.
PNG = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mP8z8BQDwAEhQGAhKmM"
    "IQAAAABJRU5ErkJggg=="
)

failures: list[str] = []
checks = 0


def check(case: str, what: str, got, want) -> None:
    global checks
    checks += 1
    if got != want:
        failures.append(f"{case}: {what}\n    got:  {got!r}\n    want: {want!r}")


def outline(doc) -> list[tuple]:
    """(tag, text, image count, page size) for every body child.

    Everything a section break carries that a reader would notice is in pgSz —
    orientation, width, height — so comparing it catches a lost break even when
    the paragraph holding it is still there.
    """
    out = []
    for el in doc.element.body:
        tag = el.tag.replace(W, "")
        text = "".join(t.text or "" for t in el.iter(f"{W}t"))
        images = sum(1 for _ in el.iter(f"{W}drawing"))
        sect = el.find(f"{W}pPr/{W}sectPr") if tag == "p" else (el if tag == "sectPr" else None)
        page = None
        if sect is not None:
            sz = sect.find(f"{W}pgSz")
            if sz is not None:
                page = (sz.get(f"{W}orient"), sz.get(f"{W}w"), sz.get(f"{W}h"))
        out.append((tag, text, images, page))
    return out


def saved(doc):
    """Round-trip through the file, as the real output does."""
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return Document(buf)


def surgery(doc):
    """What cli.run() does to the document, minus the citation marking."""
    idx = find_biblio_index(doc)
    if idx is None:
        raise AssertionError("fixture has no References heading")
    end = biblio_end_index(doc, idx)
    text = "\n".join(p.text for p in doc.paragraphs[idx + 1:end])
    entries = parse_bibliography(text)
    remove_range(doc, idx, end)
    return idx, end, entries


def picture_paragraph(doc, caption: str) -> None:
    doc.add_paragraph(caption)
    with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as fh:
        fh.write(PNG)
        path = fh.name
    try:
        doc.add_picture(path, width=Inches(1))
    finally:
        os.unlink(path)


# --- case 1: figures, tables and a landscape section after the references ----
#
# The shape of a real submitted manuscript: text, references, then the figure
# legends with their images, then a landscape page holding the wide tables.

def case_trailing_matter() -> None:
    case = "trailing figures/tables/landscape"
    doc = Document()
    doc.add_paragraph("Spinopelvic fixation is standard.1 Breach rates vary.2-3")
    doc.add_paragraph("References")
    doc.add_paragraph("1. Chang TL, Sponseller PD. Low profile pelvic fixation. Spine. 2009;34:436.")
    doc.add_paragraph("2. Kebaish KM. Sacropelvic fixation. Spine. 2010;35:2245.")
    doc.add_paragraph("3. Gertzbein SD, Robbins SE. Accuracy of pedicular screw placement. 1990;15:11.")
    doc.add_paragraph("")  # the author's spacing before the figures
    picture_paragraph(doc, "Figure 1. PRISMA 2020 flow diagram.")

    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = Inches(11.69), Inches(8.27)
    doc.add_paragraph("Tables")
    doc.add_paragraph("Table 1. Characteristics of the included studies.")
    doc.add_table(rows=2, cols=3)
    doc.add_paragraph("NR = not reported.")
    doc.add_paragraph("Table 2. Pooled breach proportion by technique.")
    doc.add_table(rows=2, cols=2)

    before = outline(doc)
    heading = next(i for i, row in enumerate(before) if row[1] == "References")
    tail_start = next(i for i, row in enumerate(before) if row[1].startswith("Figure 1."))

    idx, end, entries = surgery(doc)
    after = outline(saved(doc))

    check(case, "entries parsed", sorted(entries), [1, 2, 3])
    check(case, "last entry is not polluted by the figure legend",
          "Figure" in entries[3], False)
    check(case, "text before the bibliography", after[:heading], before[:heading])
    check(case, "everything after the bibliography", after[heading + 1:], before[tail_start:])
    check(case, "images kept", sum(r[2] for r in after), sum(r[2] for r in before))
    check(case, "tables kept", sum(1 for r in after if r[0] == "tbl"),
          sum(1 for r in before if r[0] == "tbl"))
    check(case, "no two tables left adjacent",
          [1 for a, b in zip(after, after[1:]) if a[0] == "tbl" == b[0]], [])
    check(case, "page setup of every section",
          [r[3] for r in after if r[3]], [r[3] for r in before if r[3]])
    check(case, "removal range", (idx, end), (heading, heading + 3 + 1))


# --- case 2: the section break sits on the last reference itself -------------
#
# Word puts a section break in the paragraph mark of the paragraph it follows,
# so a break placed at the end of the reference list lives inside the last
# reference. That paragraph has to go, but its break must not.

def case_break_on_last_reference() -> None:
    case = "section break inside the bibliography"
    doc = Document()
    doc.add_paragraph("Body text with a citation.1")
    doc.add_paragraph("References")
    doc.add_paragraph("1. Chang TL, Sponseller PD. Low profile pelvic fixation. Spine. 2009;34:436.")

    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = Inches(11.69), Inches(8.27)
    doc.add_paragraph("Table 1. A wide table.")
    doc.add_table(rows=2, cols=4)

    # Move the portrait break onto the reference paragraph, the way Word does.
    body = doc.element.body
    paragraphs = [el for el in body if el.tag == f"{W}p"]
    holder, reference = paragraphs[3], paragraphs[2]
    pPr = holder.find(f"{W}pPr")
    sect = pPr.find(f"{W}sectPr")
    pPr.remove(sect)
    body.remove(holder)
    ref_pPr = reference.find(f"{W}pPr")
    if ref_pPr is None:
        ref_pPr = reference.makeelement(f"{W}pPr", {})
        reference.insert(0, ref_pPr)
    ref_pPr.append(sect)

    before = outline(doc)
    idx, end, entries = surgery(doc)
    after = outline(saved(doc))

    check(case, "entries parsed", sorted(entries), [1])
    check(case, "page setup of every section",
          [r[3] for r in after if r[3]], [r[3] for r in before if r[3]])
    check(case, "the reference text is gone",
          any("Chang TL" in r[1] for r in after), False)
    # before[1:3] is the heading and the reference; everything else must survive.
    check(case, "everything outside the bibliography",
          [r[:3] for r in after if r[1]],
          [r[:3] for r in (before[:1] + before[3:]) if r[1]])


# --- case 3: an appendix after the references --------------------------------

def case_appendix() -> None:
    case = "appendix after the references"
    doc = Document()
    doc.add_paragraph("Body text.1")
    doc.add_paragraph("References")
    doc.add_paragraph("1. Kebaish KM. Sacropelvic fixation. Spine. 2010;35:2245.")
    doc.add_paragraph("Appendix")
    doc.add_paragraph("Search string: (S2AI OR sacral-alar-iliac) AND breach.")

    before = outline(doc)
    idx, end, entries = surgery(doc)
    after = outline(saved(doc))

    check(case, "entries parsed", sorted(entries), [1])
    check(case, "the appendix survives", [r[1] for r in after if r[0] == "p"],
          [r[1] for r in before[:1] + before[3:] if r[0] == "p"])


def main() -> int:
    case_trailing_matter()
    case_break_on_last_reference()
    case_appendix()

    if failures:
        print(f"FAIL — {len(failures)} of {checks} checks")
        for f in failures:
            print("  " + f)
        return 1
    print(f"PASS — {checks} checks, the bibliography and nothing else was removed")
    return 0


if __name__ == "__main__":
    sys.exit(main())
