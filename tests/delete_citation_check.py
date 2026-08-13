"""Check the review outcome that removes a citation from the sentence.

"Leave flagged" keeps an unresolved reference visible as {NEEDS REVIEW: n}.
Deleting is the other answer: the reference is in the list but the citation does
not belong here, so the marker comes out and nothing replaces it.

What makes this more than a substitution is the spacing. A marker is written
against the word before it — "knee OA 2." — so removing only the digits leaves
"knee OA ." or a doubled space. Every case below is a sentence the deletion has
to leave readable, and the mixed cases check that deleting one reference from a
group of three does not disturb the two that stay.

    python tests/delete_citation_check.py
"""
from __future__ import annotations

import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from docx import Document  # noqa: E402

from zotprep.docx_writer import make_renderer, mark_body  # noqa: E402


class Cand:
    def __init__(self, n):
        self.year, self.corporate, self.authors = 2000 + n, None, [f"author{n}"]


class Res:
    def __init__(self, status, n):
        self.status = status
        self.candidate = Cand(n) if status in ("ACCEPTED", "FROM_TEXT") else None


def run(sentence, statuses):
    """Rewrite one sentence, with each reference in the state given."""
    doc = Document()
    doc.add_paragraph(sentence)
    doc.add_paragraph("References")
    results = {n: Res(s, n) for n, s in statuses.items()}
    keys = {n: f"KEY{n:03d}" for n, s in statuses.items() if s == "ACCEPTED"}
    render = make_renderer(results, keys, "1234567", style="plain", refs=None, warnings=[])
    mark_body(doc, 1, render)
    return doc.paragraphs[0].text


# (sentence, statuses, expected text after rewriting)
CASES = [
    # the marker and the space that held it both go
    ("The burden is large 1. It has been studied elsewhere.",
     {1: "DROPPED"},
     "The burden is large. It has been studied elsewhere."),
    # mid-sentence: one space must survive, not two
    ("Reported in cohorts 4,5 and confirmed since.",
     {4: "DROPPED", 5: "DROPPED"},
     "Reported in cohorts and confirmed since."),
    # end of paragraph, with nothing after the marker to fall back on
    ("No sustained benefit was observed 9",
     {9: "DROPPED"},
     "No sustained benefit was observed"),
    # brackets go with it
    ("As reported previously [12] the effect persists.",
     {12: "DROPPED"},
     "As reported previously the effect persists."),
    # deleting one of a group leaves the others in place
    ("Three sources agree 1,2,3.",
     {1: "ACCEPTED", 2: "DROPPED", 3: "ACCEPTED"},
     "Three sources agree {Author1, (2001)}{Author3, (2003)}."),
    # a deleted reference alongside one still flagged
    ("Two sources 1,2.",
     {1: "DROPPED", 2: "REVIEW"},
     "Two sources {NEEDS REVIEW: ref 2}."),
    # nothing deleted: unchanged behaviour
    ("One source 1.",
     {1: "REVIEW"},
     "One source {NEEDS REVIEW: ref 1}."),
]


def main() -> int:
    failures = []
    for sentence, statuses, want in CASES:
        got = run(sentence, statuses)
        if got != want:
            failures.append(f"  {sentence!r}\n    want {want!r}\n    got  {got!r}")

    if failures:
        print(f"FAIL  delete: {len(failures)}/{len(CASES)} sentences wrong after deletion")
        print("\n".join(failures))
        return 1
    print(f"PASS  delete: {len(CASES)} sentences, spacing and neighbours intact")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
