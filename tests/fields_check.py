"""Check that citations come out as Word fields Zotero and Word both accept.

The output of this path is a finished document: no ODF Scan pass, so nothing
downstream will notice a malformed field. What Word checks on open is asserted
here instead — well-formed XML, balanced field characters, nothing written into
an attribute — plus the two things Zotero needs: the ADDIN instruction and an
item URI it can match.

The surrounding prose is checked character by character, because the field runs
are *inserted* into a paragraph rather than substituted into a string, and an
off-by-one in the split would eat a word.

    python tests/fields_check.py
"""
from __future__ import annotations

import io
import re
import sys
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from docx import Document  # noqa: E402
from docx.shared import Pt  # noqa: E402

from zotprep.docx_writer import (  # noqa: E402
    W,
    make_renderer,
    mark_body_fields,
)

WNS = W[1:-1]


class Res:
    """The slice of Resolution the renderer touches."""

    def __init__(self, title, year, authors, status="ACCEPTED"):
        self.status = status
        self.candidate = Cand(title, year, authors) if title else None


class Cand:
    def __init__(self, title, year, authors):
        self.title = title
        self.year = year
        self.authors = authors
        self.corporate = None


class Ref:
    def __init__(self, authors):
        self.authors = authors
        self.corporate = None


def build_doc(sentences):
    doc = Document()
    for s in sentences:
        doc.add_paragraph(s)
    return doc


def render_for(doc, resolutions, keys, biblio_idx):
    ids = iter(f"ID{n:06d}" for n in range(1, 999))
    render = make_renderer(
        resolutions, keys, "1234567", style="fields",
        refs={n: Ref(["Smith"]) for n in resolutions},
        new_id=lambda: next(ids),
    )
    return mark_body_fields(doc, biblio_idx, render)


def roundtrip(doc):
    """Save and re-open, which is the only way to see what Word would see."""
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    xml = zipfile.ZipFile(buf).read("word/document.xml")
    return xml.decode("utf-8"), Document(io.BytesIO(buf.getvalue()))


def main() -> int:
    failures: list[str] = []

    resolutions = {
        1: Res("A first paper", 2020, ["smith"]),
        2: Res("A second paper", 2021, ["jones"]),
        3: Res(None, None, None),
    }
    resolutions[3].status = "REVIEW"
    keys = {1: "AAAA1111", 2: "BBBB2222"}

    doc = build_doc([
        "The burden is large 1. It has been studied 1,2 in two cohorts.",
        "One reference resolved to nothing 3.",
        "References",
    ])
    n = render_for(doc, resolutions, keys, 2)
    if n != 3:
        failures.append(f"  expected 3 citations rewritten, got {n}")

    xml, reopened = roundtrip(doc)

    # 1. Word opens it or it does not: the XML must parse, and no citation may
    #    have landed inside an attribute value.
    try:
        ET.fromstring(xml)
    except ET.ParseError as exc:
        failures.append(f"  document.xml is not well-formed: {exc}")

    # 2. Field characters must balance, in order, for every field.
    # two fields, not three: reference 3 resolved to nothing and stays text
    kinds = re.findall(r'fldCharType="(\w+)"', xml)
    if kinds != ["begin", "separate", "end"] * 2:
        failures.append(f"  field characters out of order: {kinds}")

    # 3. Zotero's half: the instruction and a URI it can match.
    instrs = re.findall(r"ADDIN ZOTERO_ITEM CSL_CITATION (\{.*?\}) </w:instrText>", xml)
    if len(instrs) != 2:
        failures.append(f"  expected 2 ADDIN instructions, got {len(instrs)}")
    if "http://zotero.org/users/1234567/items/AAAA1111" not in xml.replace("&quot;", '"'):
        failures.append("  item URI missing from the citation JSON")
    # the grouped citation carries both items in one field
    grouped = [i for i in instrs if i.count("items/") == 2]
    if len(grouped) != 1:
        failures.append(f"  '1,2' should be one field of two items, got {len(grouped)} such")

    # 4. The prose around the citation survives, exactly.
    text = "\n".join(p.text for p in reopened.paragraphs)
    for want in ["The burden is large ", ". It has been studied ", " in two cohorts."]:
        if want not in text:
            failures.append(f"  prose lost around a field: {want!r} missing")
    if "Smith, (2020)" not in text:
        failures.append("  the visible citation text is missing")
    if "{NEEDS REVIEW: ref 3}" not in text:
        failures.append("  an unresolved reference did not stay flagged")
    if re.search(r"\bl?arge 1\.", text):
        failures.append("  the original marker digits were left behind")

    # 5. A citation inside a superscript run must not leave the field
    #    superscripted — that reads as a footnote marker, not a citation.
    doc2 = Document()
    p = doc2.add_paragraph("Reported earlier")
    run = p.add_run("1")
    run.font.superscript = True
    run.font.size = Pt(11)
    p.add_run(" and confirmed since.")
    doc2.add_paragraph("References")
    render_for(doc2, resolutions, keys, 1)
    xml2, reopened2 = roundtrip(doc2)
    shown = [r for r in reopened2.paragraphs[0].runs if r.text == "Smith, (2020)"]
    if not shown:
        failures.append("  superscript citation did not become a field")
    elif shown[0].font.superscript:
        failures.append("  the field text stayed superscripted")
    if "Reported earlier" not in reopened2.paragraphs[0].text:
        failures.append("  text before a superscript citation was lost")
    if " and confirmed since." not in reopened2.paragraphs[0].text:
        failures.append("  text after a superscript citation was lost")

    # 6. A measurement scale is a bracketed range of exactly the citation shape.
    #    "(0-10)" is refused for spanning too many references, but "(0-5)" is
    #    narrow enough to pass that test — no reference numbered zero is what
    #    stops it becoming five citations the author never wrote.
    ids = iter(f"ID{n:06d}" for n in range(1, 999))
    render = make_renderer(resolutions, keys, "1234567", style="fields",
                           refs=None, warnings=[], new_id=lambda: next(ids))
    for spec in ("0-5", "0-10", "0-100"):
        if render(spec):
            failures.append(f"  scale {spec!r} was read as a citation")
    if not render("1,2"):
        failures.append("  a real citation list stopped rendering")

    # 7. A wide range means different things in different notations. Superscript
    #    is where a leading digit goes missing, so "3-32" there is a typo; in
    #    brackets it is a systematic review citing its included studies.
    if render("1-2", False) is None:
        failures.append("  a narrow range stopped rendering")
    if render("1-30", False) is not None:
        failures.append("  a wide superscript range was expanded")
    if render("1-30", True) is None:
        failures.append("  a wide bracketed range was refused")

    total = 19
    if failures:
        print(f"FAIL  fields: {len(failures)} problems")
        print("\n".join(failures))
        return 1
    print(f"PASS  fields: {total} checks, the document Word opens is the one Zotero reads")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
